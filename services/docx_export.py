from __future__ import annotations

from io import BytesIO


def build_practice_docx(subject: str, exam_label: str, practice_text: str) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed.") from exc

    document = Document()
    document.add_heading("GCSE Practice Worksheet", level=1)
    document.add_paragraph(f"Subject: {subject}")
    document.add_paragraph(f"Paper/component: {exam_label}")
    document.add_paragraph("")

    document.add_heading("Generated Practice", level=2)
    for block in practice_text.splitlines():
        line = block.strip()
        if not line:
            document.add_paragraph("")
            continue

        if line.startswith("#"):
            document.add_heading(line.lstrip("#").strip(), level=3)
            continue

        paragraph = document.add_paragraph()
        run = paragraph.add_run(line)
        if any(symbol in line for symbol in ["=", "^", "√", "π", "Δ", "λ"]):
            run.font.name = "Consolas"
            run.font.size = Pt(10)

    document.add_page_break()
    document.add_heading("Answer Space", level=2)
    for index in range(1, 11):
        document.add_paragraph(f"Question {index}")
        for _ in range(5):
            document.add_paragraph("________________________________________________________________")

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
